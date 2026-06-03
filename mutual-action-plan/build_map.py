#!/usr/bin/env python3
"""
Build a Mixmax Mutual Action Plan (MAP) / Trial Success Plan .docx from JSON inputs.

Usage:
    python3 build_map.py --inputs map-inputs.json --out MAP-Customer-2026-06-02.docx

JSON input shape: see SKILL.md in this folder.
"""
import argparse
import json
import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)


PURPLE = RGBColor(0x7C, 0x3A, 0xED)
DARK = RGBColor(0x1A, 0x1D, 0x23)


def shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def add_para(doc, text, *, size=11, bold=False, color=None, align=None, space_after=4):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    if color is not None:
        r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)
    return p


def build(inputs: dict, out_path: str) -> None:
    doc = Document()

    # Base font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Margins
    for s in doc.sections:
        s.left_margin = Cm(2.0)
        s.right_margin = Cm(2.0)
        s.top_margin = Cm(2.0)
        s.bottom_margin = Cm(2.0)

    # ==== BRAND HEADER ====
    p = doc.add_paragraph()
    r = p.add_run("MIXMAX")
    r.font.size = Pt(14); r.bold = True; r.font.color.rgb = PURPLE

    # ==== TITLE ====
    p = doc.add_paragraph()
    r = p.add_run("Mutual Action Plan")
    r.font.size = Pt(28); r.bold = True; r.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(2)

    # Customer + date range
    add_para(doc, inputs["customer"], size=14, bold=True, space_after=2)
    add_para(doc, inputs["date_range"], size=11, color=DARK, space_after=12)

    # ==== RECIPIENTS ====
    add_para(doc, "To:", bold=True, space_after=2)
    for rec in inputs.get("recipients", []):
        name = rec.get("name", "")
        title = rec.get("title", "")
        add_para(doc, f"{name}, {title}".strip(", "), size=11, space_after=2)

    doc.add_paragraph()

    # ==== SALUTATION + OPENING ====
    salutation = inputs.get("salutation_first_names") or ""
    add_para(doc, f"Dear {salutation},", space_after=8)

    opening = doc.add_paragraph()
    opening.add_run(
        "We are excited to partner with you as you evaluate Mixmax. The purpose of this "
        "Mutual Action Plan is to outline the core activities and milestones that will help "
        "you decide whether Mixmax is the right fit against your team's specific challenges "
        "and goals. Please review and share any feedback to ensure this accurately reflects "
        "your priorities."
    ).font.size = Pt(11)

    doc.add_paragraph()

    # ==== BUSINESS CHALLENGES ====
    add_para(doc, "The business challenges your team is currently experiencing include:", space_after=4)
    for ch in inputs.get("business_challenges", []):
        add_bullet(doc, ch)

    doc.add_paragraph()

    # ==== IDEAL SOLUTION ====
    add_para(doc, "Here is a recap of what an ideal solution looks like for your team:", space_after=4)
    for outcome in inputs.get("desired_outcomes", []):
        add_bullet(doc, outcome)

    doc.add_paragraph()

    # ==== TRIAL DESCRIPTION ====
    decision_date = inputs.get("decision_date", "")
    desc = doc.add_paragraph()
    desc.add_run(
        "To evaluate whether Mixmax delivers on these outcomes, we recommend completing the "
        "following core activities during your evaluation. These will guide our check-in calls "
        f"and form the basis of your decision by {decision_date}."
    ).font.size = Pt(11)

    doc.add_paragraph()

    # ==== TRIAL ACTIVITIES ====
    add_para(doc, "Trial Activities", size=16, bold=True, color=DARK, space_after=4)
    for act in inputs.get("trial_activities", []):
        add_bullet(doc, act)

    doc.add_paragraph()

    # ==== SUCCESS METRICS ====
    add_para(doc, "Success Metrics", size=16, bold=True, color=DARK, space_after=4)
    add_para(
        doc,
        "At the end of the evaluation, we will evaluate the following leading indicators to "
        "determine whether Mixmax is the right solution for your team:",
        space_after=6,
    )

    metrics = inputs.get("success_metrics", [])
    mt = doc.add_table(rows=1, cols=2)
    mt.style = "Light Grid Accent 1"
    mt.autofit = False
    mt.columns[0].width = Cm(6)
    mt.columns[1].width = Cm(10)

    hdr = mt.rows[0].cells
    for i, txt in enumerate(["Metric", "Target / Signal"]):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(txt); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = DARK
        shade_cell(hdr[i], "F4F4F8")

    for m in metrics:
        row = mt.add_row().cells
        row[0].text = ""
        p1 = row[0].paragraphs[0]; r1 = p1.add_run(m.get("metric", "")); r1.font.size = Pt(11)
        row[1].text = ""
        p2 = row[1].paragraphs[0]; r2 = p2.add_run(m.get("target", "")); r2.font.size = Pt(11)

    doc.add_paragraph()

    # ==== ACTIVITY COMPLETION SCHEDULE ====
    add_para(doc, "Activity Completion Schedule", size=16, bold=True, color=DARK, space_after=4)

    schedule = inputs.get("schedule", [])
    st = doc.add_table(rows=1, cols=3)
    st.style = "Light Grid Accent 1"
    st.autofit = False
    st.columns[0].width = Cm(9)
    st.columns[1].width = Cm(4)
    st.columns[2].width = Cm(3)

    shdr = st.rows[0].cells
    for i, txt in enumerate(["Activity / Milestone", "Date", "Status"]):
        shdr[i].text = ""
        p = shdr[i].paragraphs[0]
        r = p.add_run(txt); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = DARK
        shade_cell(shdr[i], "F4F4F8")

    for phase_block in schedule:
        phase = phase_block.get("phase", "")
        date = phase_block.get("date", "")
        activities = phase_block.get("activities", [])
        row = st.add_row().cells
        row[0].text = ""
        p = row[0].paragraphs[0]; r = p.add_run(phase); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = DARK
        row[1].text = ""
        p2 = row[1].paragraphs[0]; r2 = p2.add_run(date); r2.font.size = Pt(11)
        row[2].text = ""
        for c in row:
            shade_cell(c, "F4F4F8")
        for act in activities:
            ar = st.add_row().cells
            ar[0].text = ""
            ap = ar[0].paragraphs[0]; ap.paragraph_format.left_indent = Cm(0.5)
            ar0 = ap.add_run("•  " + act); ar0.font.size = Pt(11)
            ar[1].text = ""
            ap2 = ar[1].paragraphs[0]; ar1 = ap2.add_run(date); ar1.font.size = Pt(11)
            ar[2].text = ""

    doc.add_paragraph()

    # ==== RESOURCES ====
    add_para(doc, "Resources to Get You Started", size=16, bold=True, color=DARK, space_after=4)
    default_resources = [
        "Mixmax Help Center — setup, integration, and feature reference",
        "Mixmax Academy — step-by-step setup by user role",
        "Mixmax How-to YouTube — visual walkthroughs",
        "Salesforce Integration Guide — field mapping and setup",
        "Deliverability & Rate Limiting — domain setup, sending limits, warm-up",
        "Sequences Best Practices — recommended cadence structures and step timing",
    ]
    for r_text in inputs.get("resources", default_resources):
        add_bullet(doc, r_text)

    doc.add_paragraph()

    # ==== CLOSING ====
    closing = doc.add_paragraph()
    closing.add_run(
        "Please feel free to reach out at any time with questions. We are committed to making "
        "this evaluation a success and look forward to reviewing results together on "
        f"{decision_date}. We will loop in your dedicated Customer Success Manager during the "
        "evaluation who will work directly with your power users to ensure a strong start and "
        "drive adoption across the full team."
    ).font.size = Pt(11)

    doc.add_paragraph()
    add_para(doc, "Sincerely,", space_after=2)
    doc.add_paragraph()

    rep = inputs.get("rep", {})
    add_para(doc, f"{rep.get('name','')}, {rep.get('title','')}".strip(", "), space_after=2)
    add_para(doc, rep.get("email", ""), space_after=2)
    add_para(doc, rep.get("phone", ""), space_after=2)

    # Trailing brand
    p = doc.add_paragraph()
    r = p.add_run("Mixmax"); r.font.size = Pt(11); r.bold = True; r.font.color.rgb = PURPLE

    # ==== SAVE ====
    os.makedirs(os.path.dirname(os.path.abspath(out_path)), exist_ok=True)
    doc.save(out_path)
    print(f"Saved: {out_path}")
    print(f"Size: {os.path.getsize(out_path)} bytes")


def main():
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--inputs", required=True, help="Path to JSON inputs file")
    ap.add_argument("--out", required=True, help="Path to output .docx")
    args = ap.parse_args()

    with open(args.inputs, "r") as f:
        data = json.load(f)

    required = ["customer", "date_range", "decision_date", "recipients", "rep"]
    missing = [k for k in required if k not in data]
    if missing:
        print(f"ERROR: missing required inputs: {missing}", file=sys.stderr)
        sys.exit(2)

    build(data, args.out)


if __name__ == "__main__":
    main()
